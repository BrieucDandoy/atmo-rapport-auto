from functools import wraps,reduce
import pandas as pd
import numpy as np
import toml
from docx import Document
from docx.shared import Pt,Inches,Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from add_float_picture import add_float_picture

import time
import configparser
import re
from sklearn.linear_model import LinearRegression
import logging
import matplotlib.pyplot as plt
#enlève le warning sur le nombre max de plot
plt.rcParams.update({'figure.max_open_warning': 0})
from io import BytesIO
logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', datefmt='%d-%b-%y %H:%M:%S')


class Rapporteur():
    """Classe permettant de générer un rapport automatiquement à partir de données générées par le programme d'extraction des données
    des microcapteurs, on utilise directement le fichier excel pour rédiger le rapport"""
    def __init__(self,nom_rapport):
        self.pas=10
        self.polluants = ["PM1","PM2.5","PM10","NO2","O3","NH3","CO","CO2","SO2","H2S","HCL","Mercaptan","COV","Alcools"]

        if nom_rapport == "":
            nom_rapport = str(f'{time.time()}.docx')
        if nom_rapport[-5:] != ".docx":
            nom_rapport = f'{nom_rapport}.docx'
        self.nom_rapport = nom_rapport
        self.__init_doc__()
        path = "config_rapport.toml"
        with open(path,'r',encoding="utf-8") as f:
            dic_config = toml.load(f)

        self.LQS = dic_config["LQ"]

        self.shadow_limit = dic_config['shadow limit']
        self.options = dic_config["document"]

        self.__init_infos_capteurs()
        
    
    def get_capteur(self,ref=False):
        """On note les noms des capteurs et de la référence dans 2 listes"""
        self.capteurs = list({self.remove_end(item) for item in self.df.columns if 'Réf' not in item and 'Date' not in item})
        if ref:
            self.references = [col for col in self.df.columns if 'Réf' in col]
            self.reference = self.remove_end(self.references[0])
        self.types_capteur = []
        for type_capteur in self.info_capteurs.keys():
            for capteur in self.capteurs:
                if type_capteur.lower() in capteur.lower():
                    self.types_capteur.append(type_capteur)
                    break

    def __init_infos_capteurs(self):
        self.info_capteurs = self.ini_to_dic('information capteurs DUMMY.ini')

    def ini_to_dic(self,ini_file):
        config = configparser.ConfigParser()
        config.optionxform = lambda option: option  # preserve case for letters
        config.read(ini_file,encoding='utf-8')
        config_dict = {}
        for section in config.sections():
            config_dict[section] = {}
            for key, value in config.items(section):
                config_dict[section][key] = value
        return config_dict


    def load_from_csv(self,nom_fichier,ref):
        self.df = pd.read_csv(nom_fichier)
        self.get_capteur(ref)

    def load_from_excel(self,file=None,ref=False):

        self.mesures_colonnes = {}
        try:
            excel = pd.ExcelFile(file)
        except Exception as e:
            logging.critical("Invalid file name\n",e)
            return None
        liste_df = []
        #on va prendre chaque feuille du fichier excel et on les met dans une liste sous forme de dataframe
        for sheet in excel.sheet_names:
            df = pd.read_excel(excel,sheet)
            liste_column = []
            for item in df.columns:
                if type(item) == int or "Unnamed:" in item:
                    break
                else:
                    liste_column.append(item)
            df = df[liste_column]
            self.mesures_colonnes[sheet] = liste_column
            df.columns = list(map(lambda x : x if 'Date'== x else f'{x}_{sheet}',df.columns))
            liste_df.append(df)
        #on fusionne les df frame sur les Dates, comme on ne peut les fusionner que 2 par 2, on utilise la fonction reduce
        dataframe = reduce(lambda x,y :(x.merge(y,on='Date')),liste_df)
        for key in self.mesures_colonnes.keys():
            self.mesures_colonnes[key].remove('Date')
            self.mesures_colonnes[key] = [f'{item}_{key}' for item in self.mesures_colonnes[key]]

        self.df = dataframe

        self.polluants = list(set(self.polluants) & set(self.mesures_colonnes.keys())) #intersection entre les liste pour différencier
        #les mesures de confort et les polluants PRESENTS dans le fichier excel
        self.get_capteur(ref)

    def save_doc(func):
        """Ce decorator sert à simplifier l'enregistrement du document"""
        @wraps(func)
        def wraper(self,*args,**kwargs):
            output = func(self,*args,**kwargs)
            self.document.save(self.nom_rapport)
            return output
        return wraper
    
    def __init_doc__(self):
        """Cette méthode sert à initialiser le document word"""
        self.document = Document()
        self.document.save(self.nom_rapport)

    def remove_margin(self,section):
        """Enlève les marges de la section dans le document"""
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0.1)
        section.right_margin = Cm(0.1)


    def remove_end(self,text)->str:
       #"""enlève tous ce qui est après le dernier _"""
        match = re.search("_[^_]*$", text)
        if match:
            fin = match.group(0)
            text = text[:-len(fin)]
        return text
    def add_title(self,titre='',level=1,center=False):
        """ajoute un titre, level règle sa taille et center l'aligne au centre ou pas"""
        h = self.document.add_heading(titre,level)
        if center:
            self.document.paragraphs[-1].alignment =  WD_ALIGN_PARAGRAPH.CENTER
        return h
  
    def add_pararagraph(self,text=''):
        """ajoute un paragraphe au document"""
        return self.document.add_paragraph(text)


    def float_to_string(self,num):
        """enlève la partie décimale si elle est nulle si le nombre vaut 0, empêche d'afficher -0"""
        str_num = str(num)
        if '.' in str_num:
            str_num = str_num.rstrip('0').rstrip('.')
        str_num = str_num.replace('.', ',')
        if str_num=='-0':
            str_num='0'
        return str_num
    
    
    def set_cell_color(self,cell,color):
        """met le fond de la couleur d'une cellule de table, color est hexadécimal et RGB"""
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'),color))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    def set_cell_font_size(self,cell,font_size):
        cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    def align_cell_center(self,cell):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    def df_to_word_table(self,df,smaller_index=True,multi_index=1,colorator=None,val_00=""):
        """Cette méthode sert à créer un tableau dans word à partir d'un dataframe pandas:
        
        -smaller_index sert à réduire le nom des colonnes qui peuvent parfois prendre de la place
        
        -multi_index est utilisé quand il y a plus d'une ligne de titre (cf documentation pandas)
        
        -val_00 est ce qui doit être écrit dans le coin en haut à gauche du tableau
        
        -colorator est une fonction qui renvoie une couleur en hexadécimale en fonction du contenu
        et de la position de la case. exemple pour mettre toutes les cases en rouge:
        colorator = lambda x,i,j : 'FF0000' """
        
        n_rows,n_cols = df.shape
        table = self.document.add_table(rows=n_rows+multi_index,cols= n_cols+1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Medium Grid 2 Accent 5'
        #On met les titres, si multiindex > 1 alors il y a plusieurs ligne pour les titre de colonnes
        if multi_index>1:
            for idx,name in enumerate(df.columns.names):
                cell = table.cell(idx,0)
                cell.text = name
                self.set_cell_font_size(table.cell(idx,0),8)
                self.set_cell_color(table.cell(idx,0),self.options["couleur_tableau_titre"])
                self.align_cell_center(cell)

            for col,names in enumerate(df.columns):
                for row,name in enumerate(names):
                    cell = table.cell(row,col+1)
                    cell.text = str(name)
                    self.set_cell_font_size(cell,8)
                    self.set_cell_color(cell,self.options["couleur_tableau_titre"])
                    self.align_cell_center(cell)
        else:
            for idx,name in enumerate(df.columns):
                if smaller_index and (name!='moyenne' or name != 'médiane'):
                    name = self.remove_end(name)[-3:]
                cell = table.cell(0,idx+1)
                cell.text = name
                self.set_cell_color(cell,self.options["couleur_tableau_titre"])
                self.set_cell_font_size(cell,8)
                self.align_cell_center(cell)
        
        cell_00 = table.cell(0,0)
        cell_00.text = val_00
        self.set_cell_font_size(cell_00,8)
        self.align_cell_center(cell_00)

        #  on entre les données, remove end sert à enlever tout ce qui est écrit après le dernier '_',
        #  parce que quand on créé le df, on ajoute le nom de chaque feuille à la fin du nom de chaque colonnes
        #  pour éviter les répétitions 
        for i in range(df.shape[0]):
            name = self.remove_end(str(df.index[i]))
            if smaller_index and (name!='moyenne' and name != 'médiane'):
                name = name[-3:]
            cell = table.cell(i+multi_index,0)
            cell.text = name
            self.set_cell_color( table.cell(i+multi_index,0),self.options["couleur_tableau_titre"])
            self.set_cell_font_size(cell,8)
            for j in range(df.shape[1]):
                cell = table.cell(i+multi_index, j+1)
                if colorator is not None:
                    color = colorator(df.values[i,j],i,j)                
                    self.set_cell_color(cell,color)
                cell.text = self.float_to_string(df.values[i,j])
                self.set_cell_font_size(cell,8)
                self.align_cell_center(cell)


    #colorator doit être une fonction qui prend en argument la valeur de la case ainsi que les coordonénes
    def get_color_percent(self,x,i,j):
        try:
            x = float(x)
        except:
            return self.options['couleur_invalid']
        if x > 80:
            return self.options['couleur_bon']
        elif x > 50:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']
    

    def get_color_pente_diag_invalid(self,x,i,j,):
        try:
            x = float(x)
        except:
            return self.options['couleur_invalid']
        if i == j:
            return self.options['couleur_invalid']
        if np.abs(1-x)<0.05:
            return self.options['couleur_bon']
        elif np.abs(1-x)<0.2:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']

    def get_color_pente(self,x,i,j,):
        try:
            x = float(x)
        except:
            return self.options['couleur_invalid']
        if np.abs(1-x)<0.05:
            return self.options['couleur_bon']
        elif np.abs(1-x)<0.2:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']  


    def get_color_intercept_diag_invalid(self,x,i,j):
        try:
            x=float(x)
        except:
            return self.options['couleur_invalid']
        if i == j:
            return self.options['couleur_invalid']
        if np.abs(x)<3:
            return self.options['couleur_bon']
        elif np.abs(x)<5:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']
    
    def get_color_intercept(self,x,i,j):
        try:
            x=float(x)
        except:
            return self.options['couleur_invalid']
        if np.abs(x)<3:
            return self.options['couleur_bon']
        elif np.abs(x)<5:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']
        

    def get_color_rsquare_diag_invalid(self,x,i,j):
        try:
            x = float(x)
        except:
            return self.options['couleur_invalid']
        if i == j:
            return self.options['couleur_invalid']
        if 0.7<x:
            return self.options['couleur_bon']
        elif 0.5<x:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais']

    def get_color_rsquare(self,x,i,j):
        try:
            x = float(x)
        except:
            return self.options['couleur_invalid']
        if 0.7<x:
            return self.options['couleur_bon']
        elif 0.5<x:
            return self.options['couleur_moyen']
        else:
            return self.options['couleur_mauvais'] 


    def df_taux(self,df):
        """Cette fonction renvoie le taux de valeur mesurée par rapport à la référence """
        df_out = pd.DataFrame()
        for polluant in self.polluants:
            #on ne garde que les valeurs qui ont une référence
            refs = [item for item in self.mesures_colonnes[polluant] if 'Réf' in item]
            if refs:
                df_tmp = df.dropna(subset=refs[0])
                lst_taux = []
                for ref in refs:
                    liste_colones = [item for item in self.mesures_colonnes[polluant] if 'Réf' not in item]
                for col in liste_colones:
                    taux = round(len(df_tmp[df_tmp[col].notna()]) / len(df_tmp)*100,1)
                    lst_taux.append(str(taux))
                df_out[polluant] = lst_taux
                df_out.index = [self.remove_end(item) for item in liste_colones]
        return df_out

    def get_df_sup_LQ(self,df):
        """Cette méthode renvoie le taux de valeur supérieur à LQ"""
        output = pd.DataFrame()
        min=None
        lst_max = []
        lst_min = []
        # keys = list(set(self.polluants) & set(self.LQS.keys()))
        liste_final = list(self.polluants) #on copie pour avoir une listee INDEPENDANTE
        for suf in self.polluants:
            lst = []
            ref = []
            for item in self.mesures_colonnes[suf]:
                if 'Réf' in item:
                    ref.append(item)
                else:
                    lst.append(item)
            lst_taux = []
            #S'il n'y a pas de référence on ne peut pas comparer à une référence
            if not ref:
                liste_final.remove(suf)
            else:
                df_tmp1 = df.dropna(subset=ref[0])#on vire les ligne où la référence n'a pas mesuré
                lst_max.append(df_tmp1[lst].max(axis=1).max(axis=0))
                for col in lst:
                    df_tmp = df_tmp1.dropna(subset=col)
                    if len(df_tmp) == 0:
                        lst_taux.append("ERREUR")
                        logging.info(f"Le capteur {col} n'a pas de mesure {suf} ayant une référence")    
                    else:
                        vec = df_tmp[df_tmp[col]>self.LQS[suf]][col]
                        if min is None or min>vec.min():
                            min = vec.min()
                        val = round(len(vec) / len(df_tmp[col]) * 100,1)
                        lst_taux.append(str(val))

                lst_min.append(min)
                min=None
                if lst_taux:
                    output[suf] = lst_taux
        output.index = [self.remove_end(item) for item in lst]
        deuxieme_row = [f"{round(item1)} -> {round(item2)}" for item1,item2 in zip(lst_min,lst_max)]
        values = [self.LQS[item] for item in liste_final]
        output.columns = pd.MultiIndex.from_arrays([output.columns,deuxieme_row,values],names=("mesure","Gamme des mesures","Limite de Quantification"))
        return output

    def get_ecarts(self,df,relatif=False):
        """Cete fonction donne l'écart moyen absolue de chaque capteur par rapport à la référence avec le
        quantiles """
        output = pd.DataFrame()

        for poll in self.polluants:
            lst=[]
            lst_ecart = []
            refs = []
            for capteur in self.mesures_colonnes[poll]:
                if 'Réf' in capteur:
                    refs.append(capteur)
                else:
                    lst.append(capteur)
            if refs:
                ref = refs[0]
                #on regarde les colonnes concernées
                for col in lst:
                    df_tmp = df[df[col].notna()]
                    df_tmp = df_tmp[df_tmp[ref].notna()]
                    if len(df_tmp) != 0:
                        if relatif:
                            #on remplace les valeurs négative par 0
                            ecart = list(np.abs(df_tmp[ref].clip(lower=0)-df_tmp[col].clip(lower=0)) / np.maximum(df_tmp[ref].clip(lower=0),df_tmp[col].clip(lower=0)))
                        else:
                            ecart = list(np.abs(df_tmp[ref].clip(lower=0)-df_tmp[col].clip(lower=0)))
                        for quant in range(5,100,self.pas):
                            ec = np.quantile(ecart,quant/100)
                            if relatif:
                                ec*=100
                            if np.isnan(ec):
                                lst_ecart.append("ERREUR")
                            else:
                                lst_ecart.append(ec)
                        output[col] = lst_ecart
                    else:
                        output[col] = ["ERREUR"]*int(100/self.pas)
                    lst_ecart = []
        output.index = [f'{i}' for i in range(5,100,self.pas)]
        return output.round(1)
    
    def display_graph_ecart(self,df):
        """affiche les graphique des ecart par rapport au quartile
        prend en argument un dataframe ayant pour colonnes les ecarts par quantile, en index
        les quantile et retourne une liste d'image. Exemeple:
        quantile capteur1 capteur2
            5%       10       12
            10%      28.7     24.9
            15%      39       40.2
            etc
            """
        df = df.select_dtypes(include=['float','int'])
        x = df.index
        liste_image = []
        for pol in self.polluants:
            if list(set(self.mesures_colonnes[pol]) & set(df.columns)):
                fig = plt.figure()
                ax = plt.subplot(111)
                for col in df.columns:
                    if pol in col:
                        ax.plot(x,df[col],label=self.remove_end(col)[-3:])
                box = ax.get_position()
                ax.set_position([box.x0, box.y0, box.width * 0.9, box.height])
                ax.legend(loc='center left', bbox_to_anchor=(1, 0.5))
                plt.title(f'Ecart par capteur en fonction du quantile pour le polluant {pol}')
                image = BytesIO()
                plt.savefig(image)
                liste_image.append(image)
        return liste_image

    def get_mean_ecart(self,df):
        """Cette méthode renvoie l'écart moyen pour par capteur par polluant"""
        polluants = ["PM10","PM2.5","O3","NO2"]
        df_out = pd.DataFrame()
        for polluant in polluants:
            lst_col = []
            lst_mean = []
            ref = ""
            for col in df.columns:
                if "Réf" in col and polluant in col:
                    ref = col
                elif polluant in col:
                    lst_col.append(col)
            for col in lst_col:
                moyenne = round(np.abs(df[ref]-df[col]).mean(),1)
                if np.isnan(moyenne):
                    lst_mean.append("ERREUR")
                else:
                    lst_mean.append(moyenne)
            df_out[polluant] = lst_mean
        df_out.index = [self.remove_end(item) for item in lst_col]
        return df_out.round(1)
    
    def get_stats(self,df):
        """fonction qui prend en entré un dataframe et ressort 3 dataframes contenant les R², les coefficients directeurs
        et les orddonnées à l'origine de chaque colonne entre elles"""
        liste_R = []
        liste_coef = []
        liste_intercept = []
        
        for j in df.columns:
            lst_R = []
            lst_C = []
            lst_I = []
            for i in df.columns:
                model = LinearRegression()
                df_nona = df[[i,j]].dropna()
                model.fit(df_nona[i].to_numpy().reshape(-1,1),df_nona[j].to_numpy().reshape(-1,1))
                lst_R.append(model.score(df_nona[i].to_numpy().reshape(-1,1),df_nona[j].to_numpy().reshape(-1,1)))
                lst_C.append(model.coef_[0][0])
                lst_I.append(model.intercept_[0])
            lst_R = lst_R
            lst_C = lst_C
            lst_I = lst_I
            liste_R.append(lst_R)
            liste_coef.append(lst_C)
            liste_intercept.append(lst_I)   
        indexes = list(df.columns)
        df_R = pd.DataFrame(liste_R,index=df.columns,columns = indexes)
        df_coef = pd.DataFrame(liste_coef,index=df.columns,columns = indexes)
        df_intercept = pd.DataFrame(liste_intercept,index=df.columns,columns = indexes)
        # Add new row to specifig index name
        df_R = pd.concat([df_R,
                          pd.DataFrame([df_R.mean()],index=['moyenne'],columns=df.columns),
                          pd.DataFrame([df_R.median()],index=['médiane'],columns=df.columns)])
        df_coef = pd.concat([df_coef,
                          pd.DataFrame([df_coef.mean()],index=['moyenne'],columns=df.columns),
                          pd.DataFrame([df_coef.median()],index=['médiane'],columns=df.columns)])
        df_intercept = pd.concat([df_intercept,
                    pd.DataFrame([df_intercept.mean()],index=['moyenne'],columns=df.columns),
                    pd.DataFrame([df_intercept.median()],index=['médiane'],columns=df.columns)])


        return (df_R.round(2),df_coef.round(2),df_intercept.round(2))  

    
    def add_introduction(self,title,text):
        self.add_pararagraph("").add_run("").add_picture(self.options['atmo_logo'])
        self.add_title(title,0)
        self.add_pararagraph(text)


    def add_info_capteurs(self,dic,idx):
        # Si le capteur est le premier à être afficher, il y a du texte au dessus de ses infos,
        # donc il faut décaler l'image vers le bas
        if idx==0:
            pos_y=Pt(360)
        else:
            pos_y=Pt(170)
        try:
            add_float_picture(self.add_pararagraph(),dic['image'],width=Inches(2.0), pos_x=Pt(400), pos_y=pos_y)
        except Exception as e:
            logging.error(f"Erreur pour l'image {dic['nom']}\nMessage d'erreur : {e}") 
        self.add_title(f"Informations sur le capteur : {dic['nom']}",0)
        self.add_title('Description du capteur',1)
        print(dic)
        description_capteur =f"""
        Dimensions : {  dic['dimensions']}
        Poids : {dic['poids']}
        Prix : {dic["prix"]}
        Décalage UTC : {dic['décalage']}"""
        self.add_pararagraph(description_capteur)
        self.add_title("Caractéristiques techniques")
        Caracteristiques_techniques = f"""  
        L'appareil est prévu pour mesurer les polluants suivants : {dic['polluants']}
        En complément, il mesure les paramètres de confort : {dic['confort']}
        Les modes de transfert des données sont : {dic['modes transfert']}\n
        Consultation des données : 
            Site internet : {dic['site']}
            Format de téléchargement des données : {dic['format']}
            API : {dic['api']}\n
        Alimentation électrique :{dic['alimentation']}
        """
        self.add_pararagraph(Caracteristiques_techniques)
    


    def add_info_kunak(self,dic,idx):
         # Si le capteur est le premier à être afficher, il y a du texte au dessus de ses infos,
        # donc il faut décaler l'image vers le bas
        if idx==0:
            pos_y=Pt(360)
        else:
            pos_y=Pt(170)
        try:
            add_float_picture(self.add_pararagraph(),dic['image'],width=Inches(2.0), pos_x=Pt(400), pos_y=pos_y)
        except Exception as e:
            logging.error(f"Erreur pour l'image {dic['nom']}\nMessage d'erreur : {e}") 
        
        self.add_title(f"Informations sur le capteur : {dic['nom']}",0)
        self.add_title('Description du capteur',1)
        description_capteur =f"""
        Dimensions : {  dic['dimensions']}
        Poids : {dic['poids']}"""
        self.add_pararagraph(description_capteur)
        self.add_title("Caractéristiques techniques")
        Caracteristiques_techniques = f"""
        L'appareil est prévu pour mesurer:
        -Gazs : {dic['gaz_mesurable']}
        -PMs : {dic['PM']}
        -confort : {dic['confort']}
        En complément, il mesure les paramètres de confort : {dic['confort']}
        Les modes de transfert des données sont : {dic['modes transfert']}\n
        Consultation des données : 
            Site internet : {dic['site']}
        Alimentation électrique :{dic['alimentation']}
        information batterie :{dic['batterie']}
        consomation : {dic['consomation']}
        """
        self.add_pararagraph(Caracteristiques_techniques)




    @save_doc
    def rapporter(self,intro=True,limite_Q=True,taux_de_fonc=True,ecart_rel=False,ecart_moyen=True,stats_entre_capteur=False,*args,**kwargs):
        if intro:
            texte = "Ce document a pour but de donner des statistiques à partir du document Excel fourni par l'extraction d'apis.\nToutes les données sont calculées à partir du fichier excel.\nCertaines informations supplémentaires sont disponibles s'il y a une station de référence."
            
            self.add_introduction(title="Introduction",text=texte)
            for idx,type_capteur in enumerate(self.types_capteur):
                if self.info_capteurs[type_capteur]['nom'] == 'kunak':
                    self.add_info_kunak(idx)
                self.add_info_capteurs(self.info_capteurs[type_capteur],idx)
                self.document.add_section()
        self.remove_margin(self.document.sections[-1])
        if taux_de_fonc:
            df_taux_de_fonc = self.df_taux(self.df)
            self.add_title("Taux de fonctionement des capteurs\n",center=True)
            self.df_to_word_table(df=df_taux_de_fonc,colorator= self.get_color_percent,val_00 = '   %   ')

        if limite_Q:
            df_limite_Q = self.get_df_sup_LQ(self.df)
            self.add_title("Taux des valeurs supérieurs à LQ\n",center=True)
            self.df_to_word_table(df_limite_Q,multi_index=3,val_00='   %   ')

        if ecart_rel:
            self.add_title("Ecarts relatifs par quantiles",center=True)
            ecarts = self.get_ecarts(self.df,relatif=True)
            liste_images = self.display_graph_ecart(ecarts)
            for polluant in self.polluants:
                liste_col = list(set(self.mesures_colonnes[polluant]) & set(ecarts.columns))
                if liste_col:
                    self.add_title(polluant,level=2)
                    self.df_to_word_table(ecarts[liste_col],val_00 = '   %   ')
            for image in liste_images:
                self.document.add_picture(image)
                self.document.paragraphs[-1].alignment =  WD_ALIGN_PARAGRAPH.CENTER
        if ecart_moyen:
            df_ecart_moyen = self.get_mean_ecart(self.df)
            self.add_title("Ecarts absolus moyens par capteur\n",center=True)
            self.df_to_word_table(df_ecart_moyen)
        
        if stats_entre_capteur:
            self.add_title("Statistiques de R², pentes et ordonnées à l'origine entre capteurs",center=True)
            
            for polluant in self.polluants:
                cols = [f'{capteur}_{polluant}' for capteur in self.capteurs]
                if len(cols)>1:
                    output = self.get_stats(self.df[cols])
                    self.add_title(f"R² {polluant}",level=2)
                    self.df_to_word_table(output[0],colorator=self.get_color_rsquare_diag_invalid)
                    self.add_title(f"Pentes {polluant}",level=2)
                    self.df_to_word_table(output[1],colorator=self.get_color_pente_diag_invalid)
                    self.add_title(f"Ordonnées aux origines {polluant}",level=2)
                    self.df_to_word_table(output[2],colorator=self.get_color_intercept_diag_invalid)


if __name__ == '__main__':
    rapporteur = Rapporteur('rapport_test')
    rapporteur.load_from_excel('H:/Produire/P_Metrologie/Correction et validation des micro-capteurs/4_Brieuc/divers/Test_Kunak um6 pour BDA - Copie.xlsx')
    print(rapporteur.df)
    rapporteur.rapporter(intro=True,limite_Q=True,taux_de_fonc=True,ecart_moyen=False,stats_entre_capteur=True,correction=True)




